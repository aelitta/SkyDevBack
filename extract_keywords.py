def run(input_params = {file_path:str}):

    import requests
    import PyPDF2  # для PDF файлов
    from docx import Document
    from dotenv import load_dotenv


    load_dotenv()

    OPEN_ROUTER_API_KEY = os.getenv('OPEN_ROUTER_API_KEY')

    # Извлечение текста из PDF (пример)
    def extract_text_from_pdf(file_path):
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text

    def extract_txt_file(file_path):
        # Чтение TXT файла
        with open(file_path, 'r', encoding='utf-8') as file:
            text_content = file.read()

        # Обрезаем текст если слишком длинный (ограничение контекста)
        max_length = 10000  # подстройте под вашу модель
        if len(text_content) > max_length:
            text_content = text_content[:max_length]

        return text_content


    def extract_docx_file(file_path):
        # Чтение DOCX файла
        doc = Document(file_path)
        text_content = ""

        # Извлекаем текст из всех параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # пропускаем пустые строки
                text_content += paragraph.text + "\n"

        return text_content

    def parse_table_from_docx(file_path):
        """Парсит таблицу из DOCX файла"""
        doc = Document(file_path)
        table = doc.tables[0]

        data = []

        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                keys = tuple(text)
                continue

            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys, text))
            data.append(row_data)

        return data



    def get_keywords(filepath):

        # Текст из файла
        if file_path[-3:] == 'pdf':
            file_text = extract_text_from_pdf(file_path)
        elif file_path[-3:] == 'txt':
            file_text = extract_txt_file(file_path)
        elif file_path[-4:] == 'docx':
            file_text = extract_docx_file(file_path)
            if file_text == '':
                file_text = parse_table_from_docx(file_path)
        # Запрос к OpenRouter
        API_KEY = OPEN_ROUTER_API_KEY
        API_URL = 'https://openrouter.ai/api/v1/chat/completions'
        headers = {
            "Authorization": f"Bearer {API_KEY}",
            "Content-Type": "application/json"
        }

        data = {
            "model": "deepseek/deepseek-chat-v3.1:free",
            "max_tokens" : 200,
            "messages": [
                {
                    "role": "user",
                    "content": f"Проанализируй этот документ и выдели ключевые слова для поиска кандидатов:\n\n{file_text}, выведи только ключевые слова через запятую"
                }
            ]
        }
        response = requests.post(API_URL, json=data, headers=headers)

        # Check if the request was successful
        if response.status_code == 200:
            print("API Response:", response.json())
        else:
            print("Failed to fetch data from API. Status Code:", response.status_code)

        result = response.json()
        return result['choices'][0]['message']['content']

    keywords = get_keywords(file_path)

    outparams['keywords'] = keywords
    outparams['def'] = 'Extract keywords from text'

    return outparams
